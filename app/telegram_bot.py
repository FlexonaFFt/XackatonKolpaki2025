import logging
import asyncio
from telegram import Update, Bot
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, CallbackQueryHandler
from sqlalchemy.orm import Session
from app.database import User
from telegram import InlineKeyboardButton, InlineKeyboardMarkup

# Configure logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Bot token - replace with your actual token
BOT_TOKEN = "7944563656:AAEzlgyqgM5t-z0THqwQCJzP67YrYJ4_C7I"
BOT_USERNAME = "hahaton"  # Replace with your bot's username

# Store for pending username verifications
pending_users = {}

# Global application instance
application = None

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a message when the command /start is issued."""
    user_id = update.effective_user.id
    
    # Проверяем, связан ли пользователь с аккаунтом
    from app.database import get_db, User
    db = next(get_db())
    
    # Ищем пользователя с таким telegram_id
    user = db.query(User).filter(User.telegram_id == str(user_id)).first()
    
    if user:
        # Пользователь уже связан с аккаунтом
        await update.message.reply_text(
            f"Вы уже авторизованы как {user.username}! "
            f"Вы будете получать уведомления через этого бота."
        )
    else:
        # Пользователь еще не связан с аккаунтом
        await update.message.reply_text(
            "Привет! Я бот для интеграции с вашим аккаунтом. "
            "Пожалуйста, отправьте мне ваш username с сайта, чтобы я мог связать ваши аккаунты."
        )

async def handle_username(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Handle username messages and automatically link accounts."""
    username = update.message.text.strip()
    user_id = update.effective_user.id
    
    # Get database session
    from app.database import get_db
    db = next(get_db())
    
    # Try to link the user directly
    success = verify_and_link_user(db, username, str(user_id))
    
    if success:
        await update.message.reply_text(
            f"Отлично! Ваш аккаунт с username '{username}' успешно связан с Telegram. "
            f"Теперь вы будете получать уведомления через бота."
        )
    else:
        # Store the username and user_id for later verification if needed
        pending_users[user_id] = username
        
        await update.message.reply_text(
            f"К сожалению, не удалось найти пользователя с username '{username}'. "
            f"Пожалуйста, проверьте правильность написания или зарегистрируйтесь на сайте."
        )

def get_bot_link():
    """Return the link to the bot."""
    return f"https://t.me/{BOT_USERNAME}"

async def send_message_to_user(telegram_id: str, message: str):
    """Send a message to a user by their telegram_id."""
    try:
        bot = Bot(token=BOT_TOKEN)
        await bot.send_message(chat_id=telegram_id, text=message)
        return True
    except Exception as e:
        logger.error(f"Error sending message to {telegram_id}: {e}")
        return False

async def send_format_options(telegram_id: str, post_id: int, message: str):
    """Send a message with format options buttons."""
    try:
        bot = Bot(token=BOT_TOKEN)
        
        # Создаем клавиатуру с кнопками выбора формата
        keyboard = [
            [
                InlineKeyboardButton("PDF", callback_data=f"format_pdf_{post_id}"),
                InlineKeyboardButton("DOCX", callback_data=f"format_docx_{post_id}")
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        # Отправляем сообщение с клавиатурой
        await bot.send_message(
            chat_id=telegram_id,
            text=message,
            reply_markup=reply_markup
        )
        return True
    except Exception as e:
        logger.error(f"Error sending format options to {telegram_id}: {e}")
        return False

async def handle_callback_query(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle callback queries from inline keyboards."""
    query = update.callback_query
    await query.answer()
    
    # Получаем данные из callback_data
    callback_data = query.data
    
    if callback_data.startswith("format_"):
        # Разбираем данные формата и ID поста
        parts = callback_data.split("_")
        if len(parts) == 3:
            format_type = parts[1]  # pdf или docx
            post_id = int(parts[2])
            
            # Сообщаем пользователю, что документ готовится
            await query.edit_message_text(
                text=f"Подготовка документа в формате {format_type.upper()}... Пожалуйста, подождите."
            )
            
            # Здесь будет логика создания и отправки документа
            # Пока просто отправляем сообщение о том, что функция в разработке
            await send_document_to_user(query.from_user.id, post_id, format_type)

async def send_document_to_user(telegram_id, post_id, format_type):
    """Send document to user based on post_id and format_type."""
    try:
        bot = Bot(token=BOT_TOKEN)
        
        # Получаем данные поста из базы данных
        from app.database import get_db, Post
        db = next(get_db())
        post = db.query(Post).filter(Post.id == post_id).first()
        
        if not post:
            await bot.send_message(
                chat_id=telegram_id,
                text="Извините, пост не найден."
            )
            return False
        
        # Создаем временный файл для документа
        import tempfile
        import os
        
        # Генерируем имя файла на основе заголовка поста
        safe_title = "".join(c for c in post.title if c.isalnum() or c in [' ', '_']).strip()
        safe_title = safe_title.replace(' ', '_')
        
        if format_type == "pdf":
            # Создаем PDF документ
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_file.close()
            
            doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Создаем содержимое документа
            content = []
            
            # Добавляем заголовок
            title_style = styles['Title']
            content.append(Paragraph(post.title, title_style))
            content.append(Spacer(1, 12))
            
            # Добавляем основной текст
            normal_style = styles['Normal']
            for paragraph in post.content.split('\n\n'):
                if paragraph.strip():
                    content.append(Paragraph(paragraph, normal_style))
                    content.append(Spacer(1, 6))
            
            # Создаем PDF
            doc.build(content)
            
            # Отправляем файл
            with open(temp_file.name, 'rb') as file:
                await bot.send_document(
                    chat_id=telegram_id,
                    document=file,
                    filename=f"{safe_title}.pdf",
                    caption=f"Документ: {post.title}"
                )
            
            # Удаляем временный файл
            os.unlink(temp_file.name)
            
        elif format_type == "docx":
            # Создаем DOCX документ
            from docx import Document
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_file.close()
            
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading(post.title, 0)
            
            # Добавляем основной текст
            for paragraph in post.content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Сохраняем документ
            doc.save(temp_file.name)
            
            # Отправляем файл
            with open(temp_file.name, 'rb') as file:
                await bot.send_document(
                    chat_id=telegram_id,
                    document=file,
                    filename=f"{safe_title}.docx",
                    caption=f"Документ: {post.title}"
                )
            
            # Удаляем временный файл
            os.unlink(temp_file.name)
        
        return True
    except Exception as e:
        logger.error(f"Error sending document to {telegram_id}: {e}")
        # Отправляем сообщение об ошибке пользователю
        try:
            await bot.send_message(
                chat_id=telegram_id,
                text=f"Произошла ошибка при создании документа: {str(e)}"
            )
        except:
            pass
        return False

def verify_and_link_user(db: Session, username: str, telegram_id: str):
    """Verify and link a user's account with their Telegram ID."""
    user = db.query(User).filter(User.username == username).first()
    if not user:
        return False
    
    user.telegram_id = telegram_id
    db.commit()
    
    # Log successful linking
    logger.info(f"User {username} linked with Telegram ID {telegram_id}")
    return True

async def setup_bot():
    """Set up the Telegram bot application."""
    global application
    application = Application.builder().token(BOT_TOKEN).build()
    
    # Add handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_username))
    application.add_handler(CallbackQueryHandler(handle_callback_query))  # Добавляем обработчик callback_query
    
    return application

async def start_bot():
    """Start the Telegram bot in the background."""
    global application
    
    if application is None:
        application = await setup_bot()
    
    # Start the bot without blocking
    await application.initialize()
    await application.start()
    
    # Start polling in the background
    asyncio.create_task(application.updater.start_polling())
    
    logger.info("Telegram bot started")

async def stop_bot():
    """Stop the Telegram bot."""
    global application
    
    if application:
        await application.stop()
        await application.shutdown()
        logger.info("Telegram bot shutdown complete")
